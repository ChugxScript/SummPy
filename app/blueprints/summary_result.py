from flask import Blueprint, render_template, request, session, current_app, flash, redirect, url_for, jsonify
import os 
import shutil
import fitz
from fpdf import FPDF
from datetime import datetime
from .BART.summarize_doc import SummPy
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

summary_result = Blueprint('summary_result', __name__)

def get_summarized_files():
    summarized_folder = current_app.config['SUMMARIZED_FOLDER']
    if os.path.exists(summarized_folder):
        return os.listdir(summarized_folder)
    return []

def get_uploaded_files():
    upload_folder = current_app.config['UPLOAD_FOLDER']
    if os.path.exists(upload_folder):
        return os.listdir(upload_folder)
    return []

def generate_pdf_with_first_page(summary_list, filename, original_file_path): 
    pdf = FPDF()
    
    doc = fitz.open(original_file_path)
    first_page = doc[0]
    blocks = first_page.get_text("dict")["blocks"] 

    formatted_lines = []
    previous_bottom = 0  
    
    for block in blocks:
        if "lines" not in block:
            continue
        
        for line in block["lines"]:
            for span in line["spans"]:
                current_top = span["bbox"][1]  
                font_size = span["size"]     
                
                if current_top - previous_bottom > font_size * 0.8:
                    formatted_lines.append(" ")  
              
                formatted_lines.append(span["text"])
                previous_bottom = span["bbox"][3]  


    formatted_text = "\n".join(formatted_lines)

    pdf.add_page()
    pdf.set_font("Times", size=12)
    
    for line in formatted_lines:
        if line.strip() == "":
            pdf.ln(6)  
        else:
            pdf.multi_cell(0, 7, line, align='C')
    
    section_titles = ["INTRODUCTION", "METHOD", "RESULTS", "DISCUSSION"]

    for i, summary in enumerate(summary_list):
        # Add a new page before each section (including the first one)
        pdf.add_page()

        # Set the title dynamically based on the section
        pdf.set_font("Arial", style='B', size=12)
        pdf.cell(0, 10, section_titles[i], ln=True, align='C')

        # Add the summary for the current section
        pdf.set_font("Arial", size=12)
        encoded_summary = summary.encode('latin-1', 'replace').decode('latin-1')

        # Add multi_cell for summary text to ensure newlines are preserved
        pdf.multi_cell(0, 10, encoded_summary)
        pdf.ln(10)  # Add line break after the summary

    # Ensure the folder exists before saving
    summarized_folder = current_app.config['SUMMARIZED_FOLDER']
    if summarized_folder and not os.path.exists(summarized_folder):
        os.makedirs(summarized_folder)

    pdf_path = os.path.join(summarized_folder, filename)
    print("making filename PDF: "+filename)
    pdf.output(pdf_path)

def generate_docx_with_first_page(summary_list, filename, original_file_path): 
    doc = fitz.open(original_file_path)
    first_page = doc[0]
    blocks = first_page.get_text("dict")["blocks"] 

    formatted_lines = []
    previous_bottom = 0  

    for block in blocks:
        if "lines" not in block:
            continue
        
        for line in block["lines"]:
            for span in line["spans"]:
                current_top = span["bbox"][1]  
                font_size = span["size"]     
                
                if current_top - previous_bottom > font_size * 0.8:
                    formatted_lines.append("")  # Add a blank line for spacing
              
                formatted_lines.append(span["text"])
                previous_bottom = span["bbox"][3]  

    # Create a new Word document
    docx_document = Document()
    formatted_text = "\n".join(formatted_lines)

    # Add the first page text
    docx_document.add_heading('First Page Text', level=1)
    for line in formatted_lines:
        if line.strip() == "":
            docx_document.add_paragraph()  # Add a blank line
        else:
            paragraph = docx_document.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the text

    section_titles = ["INTRODUCTION", "METHOD", "RESULTS", "DISCUSSION"]

    for i, summary in enumerate(summary_list):
        # Add a heading for each section - no need for page break before first section
        if i > 0:
            docx_document.add_page_break()
        else:
            # Add some space between first page content and first section
            docx_document.add_paragraph()
            docx_document.add_paragraph()
            
        docx_document.add_heading(section_titles[i], level=1)

        # Add the summary text
        docx_document.add_paragraph(summary)

    # Ensure the folder exists before saving
    summarized_folder = current_app.config['SUMMARIZED_FOLDER']
    if summarized_folder and not os.path.exists(summarized_folder):
        os.makedirs(summarized_folder)

    docx_path = os.path.join(summarized_folder, filename)
    print("making filename DOCX: " + filename)
    docx_document.save(docx_path)

def save_documents(uploaded_files, summarized_files):
    result_orig_folder = current_app.config['ADMIN_RAW_DOCU_FOLDER']
    result_summarized_folder = current_app.config['ADMIN_SUMM_DOCU_FOLDER']

    for i, file_name in enumerate(uploaded_files):
        # Copy original files to result_orig_folder
        source_path = os.path.join(current_app.config['UPLOAD_FOLDER'], file_name)
        dest_path = os.path.join(result_orig_folder, file_name)
        if os.path.exists(source_path):
            shutil.copy(source_path, dest_path)
        else:
            print(f"Warning: Source file {source_path} not found.")
    
    for i, file_name in enumerate(summarized_files):
        # Copy summarized files to result_summarized_folder
        summarized_file_path = os.path.join(current_app.config['SUMMARIZED_FOLDER'], file_name)
        summarized_dest_path = os.path.join(result_summarized_folder, file_name)
        if os.path.exists(summarized_file_path):
            shutil.copy(summarized_file_path, summarized_dest_path)
        else:
            print(f"Warning: Summarized file {summarized_file_path} not found.")

@summary_result.route('/summary_result', methods=['GET'])
def summary_result_page():
    uploaded_files = get_uploaded_files()
    print(f"UPLOAD_FOLDER is: {uploaded_files}")
    # 13 23 51 85 127 131

    #  get the IMRAD summary
    summarizer = SummPy()
    # results = summarizer.generate_summaries()
    results = summarizer.process_all_files()

    # make pdfs with the first page of the original PDF
    for i, result in enumerate(results):
        try:
            if i >= len(uploaded_files):
                print(f"Warning: More results ({len(results)}) than uploaded files ({len(uploaded_files)})")
                continue
                
            if not result or 'summaries' not in result or not result['summaries']:
                print(f"Warning: No summaries found for file {uploaded_files[i]}")
                continue
                
            original_file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], uploaded_files[i])
            pdf_filename = f"summary_{uploaded_files[i]}"
            docx_filename = f"summary_{uploaded_files[i].replace('.pdf', '.docx')}"
            print("Processing file:", docx_filename)
            
            # Ensure we have valid summaries
            if isinstance(result, list) and len(result) > 0:
                summaries = result
            elif isinstance(result, dict) and 'summaries' in result and len(result['summaries']) > 0:
                summaries = result['summaries']
            else:
                print(f"Warning: Invalid result format for file {uploaded_files[i]}")
                continue
                
            generate_pdf_with_first_page(summaries, pdf_filename, original_file_path)
            generate_docx_with_first_page(summaries, docx_filename, original_file_path)
        except Exception as e:
            # Log the error and continue with the next file
            print(f"Error processing file {uploaded_files[i]}: {e}")
            
            # Define log directory and ensure it exists
            log_path = os.path.join(os.getcwd(), 'app/logs')
            if not os.path.exists(log_path):
                os.makedirs(log_path)
            
            # Create log file name based on the current date with .txt extension
            log_filename = f"error_{datetime.now().strftime('%Y-%m-%d')}.txt"
            log_file_path = os.path.join(log_path, log_filename)
            
            # Append the error to the log file
            with open(log_file_path, 'a') as log_file:
                log_file.write(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Error processing file {uploaded_files[i]}: {e}\n")
            
            continue

    summarized_folder = get_summarized_files()
    save_documents(uploaded_files, summarized_folder)
    print(f"UPLOAD_FOLDER is: {uploaded_files}")
    print(f"summarized_folder is: {summarized_folder}")

    return render_template('summary_result.html', summarized_folder=summarized_folder, uploaded_files=uploaded_files)
